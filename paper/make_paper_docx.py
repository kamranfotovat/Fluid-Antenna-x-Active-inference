"""
Generate an IEEE two-column .docx first draft from our verified results.

No LaTeX/pandoc on this machine, so we build the IEEE style directly with python-docx:
Times New Roman, 24pt centered title, author block, single-column abstract, then a
continuous section break to a two-column body, numbered equations, embedded figures.

Output: paper/Active_Inference_FAS_draft.docx
Run:    python paper/make_paper_docx.py
"""

from __future__ import annotations

import os
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

HERE = os.path.dirname(os.path.abspath(__file__))
FIG = os.path.join(HERE, "..", "figures")
OUT = os.path.join(HERE, "Active_Inference_FAS_draft.docx")

TNR = "Times New Roman"


def set_cols(section, n):
    sectPr = section._sectPr
    cols = sectPr.find(qn("w:cols"))
    if cols is None:
        cols = OxmlElement("w:cols"); sectPr.append(cols)
    cols.set(qn("w:num"), str(n))
    cols.set(qn("w:space"), "360")   # ~0.25in gutter


def base_font(doc):
    st = doc.styles["Normal"]
    st.font.name = TNR
    st.font.size = Pt(10)
    st.element.rPr.rFonts.set(qn("w:eastAsia"), TNR)
    pf = st.paragraph_format
    pf.space_after = Pt(0); pf.line_spacing = 1.0


def para(doc, text="", size=10, bold=False, italic=False, align=None, after=0, before=0,
         first_indent=0.2):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    pf = p.paragraph_format
    pf.space_after = Pt(after); pf.space_before = Pt(before)
    if first_indent:
        pf.first_line_indent = Inches(first_indent)
    r = p.add_run(text)
    r.font.name = TNR; r.font.size = Pt(size); r.bold = bold; r.italic = italic
    return p, r


def heading(doc, num, title):
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(8); p.paragraph_format.space_after = Pt(4)
    r = p.add_run(f"{num}. {title.upper()}")
    r.font.name = TNR; r.font.size = Pt(10); r.font.small_caps = True; r.bold = False
    return p


def subhead(doc, letter, title):
    p = doc.add_paragraph(); p.paragraph_format.space_before = Pt(4)
    r = p.add_run(f"{letter}. "); r.font.name = TNR; r.font.size = Pt(10); r.italic = True
    r2 = p.add_run(title); r2.font.name = TNR; r2.font.size = Pt(10); r2.italic = True
    return p


def body(doc, text, indent=0.2):
    return para(doc, text, size=10, align=WD_ALIGN_PARAGRAPH.JUSTIFY, first_indent=indent)


def equation(doc, eq, num):
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(3); p.paragraph_format.space_after = Pt(3)
    r = p.add_run(eq + f"  ({num})"); r.font.name = TNR; r.font.size = Pt(10); r.italic = True
    return p


def figure(doc, fname, caption, width=3.3):
    path = os.path.join(FIG, fname)
    if os.path.exists(path):
        p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run().add_picture(path, width=Inches(width))
    p2 = doc.add_paragraph(); p2.paragraph_format.space_before = Pt(2); p2.paragraph_format.space_after = Pt(6)
    r = p2.add_run(caption); r.font.name = TNR; r.font.size = Pt(8)


def build():
    doc = Document()
    base_font(doc)
    sec = doc.sections[0]
    sec.top_margin = Inches(0.75); sec.bottom_margin = Inches(1.0)
    sec.left_margin = Inches(0.62); sec.right_margin = Inches(0.62)
    set_cols(sec, 1)

    # ---- Title + authors (single column) ----
    para(doc, "Active Inference for Port Selection in Fluid Antenna Systems Under Partial CSI",
         size=20, bold=False, align=WD_ALIGN_PARAGRAPH.CENTER, after=8, first_indent=0)
    para(doc, "Kian Fotovat¹ and Kamran Fotovat²", size=11,
         align=WD_ALIGN_PARAGRAPH.CENTER, after=2, first_indent=0)
    para(doc, "¹Dept. of Electrical and Computer Engineering, University of Tehran, Tehran, Iran",
         size=9, italic=True, align=WD_ALIGN_PARAGRAPH.CENTER, after=0, first_indent=0)
    para(doc, "²Iran University of Science and Technology, Tehran, Iran",
         size=9, italic=True, align=WD_ALIGN_PARAGRAPH.CENTER, after=8, first_indent=0)

    # ---- Abstract + index terms (single column) ----
    ab = doc.add_paragraph(); ab.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    ab.paragraph_format.left_indent = Inches(0.3); ab.paragraph_format.right_indent = Inches(0.3)
    r = ab.add_run("Abstract—"); r.font.name = TNR; r.font.size = Pt(9); r.bold = True; r.italic = True
    r2 = ab.add_run(
        "Fluid antenna systems (FAS) enhance spatial diversity by activating a few radiating ports "
        "among many candidate positions, but selecting good ports conventionally requires channel state "
        "information (CSI) across all ports, incurring prohibitive acquisition overhead. We recast FAS "
        "port selection as active inference under partial CSI: the transmitter measures only the ports it "
        "activates, infers the remaining ports from a spatio-temporal generative model (Jakes spatial "
        "correlation and a first-order temporal model) through a complex Kalman belief, and selects ports "
        "by minimizing the expected free energy (EFE), which unifies expected achievable rate, information "
        "gain, and antenna-switching cost in a single objective. A greedy submodular selector yields low "
        "decision latency, and an observe-then-precode protocol keeps the served-port CSI fresh. "
        "Simulations show that the proposed method attains 84–89% of a full-CSI genie's sum rate while "
        "observing only 20% of the ports, surpasses naive, deep-reinforcement-learning, and bandit "
        "baselines on the switching-aware objective, requires no training, and remains robust to Doppler "
        "and to model mismatch, for which the spatial correlation can be learned online.")
    r2.font.name = TNR; r2.font.size = Pt(9); r2.italic = True

    it = doc.add_paragraph(); it.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    it.paragraph_format.left_indent = Inches(0.3); it.paragraph_format.right_indent = Inches(0.3)
    it.paragraph_format.space_before = Pt(4); it.paragraph_format.space_after = Pt(6)
    r = it.add_run("Index Terms—"); r.font.name = TNR; r.font.size = Pt(9); r.bold = True; r.italic = True
    r2 = it.add_run("Fluid antenna systems, port selection, active inference, expected free energy, "
                    "partial CSI, Kalman filtering.")
    r2.font.name = TNR; r2.font.size = Pt(9); r2.italic = True

    # ---- switch to two columns ----
    doc.add_section(WD_SECTION.CONTINUOUS)
    set_cols(doc.sections[-1], 2)

    # ===================== I. Introduction =====================
    heading(doc, "I", "Introduction")
    body(doc,
         "Fluid antenna systems (FAS), also referred to as movable-antenna systems, employ radiating "
         "elements whose position can be reconfigured within a small aperture, providing spatial degrees "
         "of freedom beyond those of fixed antenna arrays [1]. By activating a small number of ports out "
         "of many candidate positions, an FAS can capture favorable channel conditions with few radio-"
         "frequency chains. The central control problem is port selection: deciding which ports to "
         "activate so as to maximize throughput.", indent=0)
    body(doc,
         "Selecting ports well, however, conventionally presupposes channel state information (CSI) over "
         "all candidate ports. Because the number of ports is large and the aperture is densely sampled, "
         "acquiring this CSI incurs substantial pilot and training overhead, which grows with the array "
         "size. Existing approaches largely sidestep rather than confront this cost. Learning-based "
         "selectors, whether deep reinforcement learning (DRL) with attention over ports [2] or online "
         "bandit methods [3], typically assume full CSI and/or require extensive training. Channel "
         "estimation methods for FAS, such as the successive Bayesian reconstructor [4] and sequential "
         "linear MMSE estimation [5], reconstruct the channel accurately but address estimation rather "
         "than the joint decision of what to sense, what to serve, and when to move.")
    body(doc,
         "We argue that port selection is most naturally posed as active acquisition under partial CSI, "
         "and that active inference [6] provides a principled framework for it. Our contributions are:")
    body(doc,
         "1) We formulate FAS port selection as active inference under partial observability, removing "
         "the unrealistic full-CSI assumption of prior work.", indent=0)
    body(doc,
         "2) We develop a model-based agent that maintains a complex Kalman belief over all ports from a "
         "spatio-temporal prior, and selects ports by minimizing an expected free energy (EFE) that "
         "unifies expected rate, information gain, and switching cost in a single objective.", indent=0)
    body(doc,
         "3) We give a greedy submodular selector with an (1−1/e) guarantee on the information term "
         "and O(NM) latency, together with an observe-then-precode protocol that makes performance robust "
         "to Doppler.", indent=0)
    body(doc,
         "4) We show through simulation that the method attains 84–89% of a full-CSI genie while "
         "observing 20% of ports, beats naive, DRL, and bandit baselines on the switching-aware "
         "objective with zero training, and can learn the spatial correlation online under model "
         "mismatch.", indent=0)

    # ===================== II. System Model =====================
    heading(doc, "II", "System Model and Problem Formulation")
    body(doc,
         "We consider a base station whose fluid antenna comprises N = Nₓ×Nᵧ candidate "
         "ports arranged on a two-dimensional grid over a sub-half-wavelength aperture. Each slot the "
         "base station activates M of the N ports (one per RF chain) to serve K single-antenna users, "
         "with N ≥ M ≥ K. Let hₖ(t) ∈ ℂᴺ denote the channel of user k over "
         "all N ports in slot t.", indent=0)
    subhead(doc, "A", "Spatio-temporal channel model")
    body(doc,
         "Ports on a dense aperture are spatially correlated. Following Jakes' model, the correlation "
         "between ports i and j is")
    equation(doc, "Rᵢⱼ = J₀(2π dᵢⱼ / λ),", 1)
    body(doc, "where J₀(·) is the zeroth-order Bessel function and dᵢⱼ the inter-port "
              "distance. The channel evolves as a first-order Gauss–Markov (AR(1)) process,")
    equation(doc, "hₖ(t) = ρ hₖ(t−1) + √(1−ρ²) eₖ(t),  "
                  "eₖ ~ ᶜᴺ(0, βₖ R),", 2)
    body(doc, "with temporal correlation ρ = J₀(2π f_D T_s) set by the Doppler f_D. The "
              "stationary distribution is hₖ ~ ᶜᴺ(0, βₖ R).")
    subhead(doc, "B", "Partial observability")
    body(doc,
         "The key departure from prior work is that full CSI is unavailable. Activating a port set S "
         "yields noisy pilots only on those M ports,")
    equation(doc, "yₖ(t) = P_S hₖ(t) + n,  n ~ ᶜᴺ(0, σₑ² I),", 3)
    body(doc, "where P_S selects the active rows; the remaining N−M ports are hidden and must be "
              "inferred.")
    subhead(doc, "C", "Precoding, rate, and objective")
    body(doc,
         "The base station applies a robust MMSE precoder W built from its belief and serves the users; "
         "the achievable rate of user k is Rₖ = log₂(1 + SINRₖ). Moving the fluid element "
         "between ports costs delay and energy, modeled by a switching cost C_t = |S_t △ S_{t−1}| "
         "and E_sw = e_sw C_t. The long-term objective is")
    equation(doc, "max Σ_t [ Σₖ Rₖ(t) − η_sw E_sw(t) ].", 4)

    # ===================== III. Method =====================
    heading(doc, "III", "Active-Inference Port Selection")
    subhead(doc, "A", "Belief (perception)")
    body(doc,
         "The agent maintains a per-user Gaussian belief q(hₖ) = ᶜᴺ(μₖ, "
         "Σₖ) over all N ports, updated by a complex Kalman filter. The predict step encodes "
         "CSI aging, μ ← ρμ and Σ ← ρ²Σ + (1−ρ²)"
         "βₖR, and the measurement update uses the Joseph form for numerical stability. Because "
         "the innovation covariance is regularized by σₑ²I, the (rank-deficient) prior "
         "requires no artificial jitter, and the belief is calibrated: its covariance matches the "
         "realized estimation error.", indent=0)
    subhead(doc, "B", "Expected free energy (action)")
    body(doc, "A candidate port set S is scored by its expected free energy (in bits):")
    equation(doc, "G(S) = −α·Prag(S) − β·Epis(S) + Switch(S),", 5)
    body(doc,
         "where the pragmatic term is the expected robust-MMSE sum rate under the belief (an imperfect-"
         "CSI lower bound that grows conservative as uncertainty grows); the epistemic term is the mutual "
         "information Epis(S) = Σₖ log₂ det(I + P_S Σₖ P_Sᴴ / σₑ"
         "²), which is monotone submodular and, through the off-diagonals of Σ, also sharpens "
         "correlated neighbors; and the switching term is η_sw e_sw |S △ S_prev|. Exploration "
         "and exploitation are thus unified in one objective rather than bolted on.")
    subhead(doc, "C", "Observe-then-precode and greedy selection")
    body(doc,
         "Two design choices make the agent both accurate and fast. First, within a slot the agent "
         "selects on the predicted belief but then senses the activated ports and precodes on the updated "
         "belief; this observe-then-precode ordering keeps the served-port CSI error near σₑ"
         "² instead of the aging floor, and, as shown below, removes the dependence on Doppler. "
         "Second, since choosing M of N is combinatorial, S is built greedily by adding the port of "
         "largest marginal G; the epistemic term's submodularity yields an (1−1/e) guarantee and the "
         "cost is O(NM), about 20 ms per slot versus seconds for exhaustive search.", indent=0)
    body(doc,
         "When the spatial correlation is unknown or non-Jakes, R can be estimated online from the "
         "co-observed pilots (a normalized empirical correlation), which restores performance under "
         "model mismatch.")

    # ===================== IV. Results =====================
    heading(doc, "IV", "Simulation Results")
    body(doc,
         "Unless stated otherwise, N = 25 (5×5, sub-λ/2), K = 3, M = 5 (a 20% observation "
         "budget), SNR 15 dB, σₑ² = 10⁻³, ρ = 0.9, and η_sw = 1, "
         "averaged over Monte-Carlo channel realizations. Baselines are a full-CSI genie (rate ceiling), "
         "a naive no-inference selector, random selection, a DRL policy (a Transformer port scorer "
         "trained by policy gradient), and a combinatorial-UCB bandit.", indent=0)
    figure(doc, "figA_observation_budget.png",
           "Fig. 1. Rate and switching-aware objective versus observation budget M/N. The proposed method "
           "tracks 73–91% of the genie rate and exceeds it on the objective for budgets ≥ ~18%.")
    body(doc,
         "Observation budget. As the budget grows, the agent tracks 73–91% of the genie's rate and, "
         "on the switching-aware objective, exceeds the genie because it barely moves the antenna. At the "
         "20% operating point it attains about 84% of genie rate.", indent=0)
    figure(doc, "figC_protocol_doppler.png",
           "Fig. 2. Observe-then-precode versus predict-then-act across Doppler. The proposed protocol "
           "holds at 84–89% of genie and is flat in ρ; predict-then-act collapses.")
    body(doc,
         "Protocol and Doppler robustness. Observe-then-precode reaches 84–89% of the genie and is "
         "essentially flat across ρ, whereas precoding on the predicted (aged) belief degrades from "
         "64% to 25% as the channel speeds up.", indent=0)
    figure(doc, "figF_pareto_frontier.png",
           "Fig. 3. Sweeping the exploration weight traces a Pareto frontier that dominates the genie: "
           "from 84% rate at near-zero switching to 89% rate, the whole frontier beats the genie "
           "objective.")
    body(doc,
         "Frontier. Sweeping the exploration weight yields a frontier that dominates the genie, from 84% "
         "rate at near-zero switching to a maximum of 89% rate; the entire frontier beats the genie's "
         "objective while switching far less.", indent=0)
    figure(doc, "figB_drl_sample_efficiency.png",
           "Fig. 4. Versus DRL. The proposed method (zero training, 20% CSI) meets or exceeds a fully "
           "trained full-CSI DRL on the objective; the DRL needs ~100–200 episodes to plateau.")
    figure(doc, "figG_bandit_comparison.png",
           "Fig. 5. Versus a bandit. The proposed method beats a combinatorial-UCB bandit at every "
           "exploration setting: comparable rate, higher objective, and roughly 100× less switching.")
    body(doc,
         "Comparison with learning baselines. The proposed agent needs no training yet meets or exceeds a "
         "fully trained full-CSI DRL on the objective, and beats the bandit at every exploration setting "
         "(+12% over its best) with about 100× less switching, because a model-free learner has no "
         "persistent per-port signal to exploit in the equal-average-power regime and thus pays a "
         "perpetual exploration cost.", indent=0)
    figure(doc, "figE_learning_mismatch.png",
           "Fig. 6. Model mismatch. On a non-Jakes (exponential-correlation) channel, learning R online "
           "recovers the oracle objective, whereas assuming Jakes loses.")
    body(doc,
         "Model mismatch. When the true correlation is non-Jakes, an agent that wrongly assumes Jakes "
         "loses performance, while learning R online recovers the oracle result.", indent=0)

    # ===================== V. Conclusion =====================
    heading(doc, "V", "Conclusion")
    body(doc,
         "We posed fluid-antenna port selection as active inference under partial CSI, unifying expected "
         "rate, information gain, and switching cost in a single expected-free-energy objective over a "
         "spatio-temporal Kalman belief. The method matches most of the full-CSI performance while "
         "observing a fraction of ports, wins the switching-aware objective against naive, DRL, and "
         "bandit baselines without any training, and is robust to Doppler and to model mismatch. Future "
         "work includes tracking a moving spatial hotspot with a latent-envelope belief, scaling to large "
         "arrays where the low-rank channel makes acquisition even cheaper, a Gaussian-process belief "
         "front-end, and a distance-weighted switching model.", indent=0)

    # ---- Acknowledgment (AI disclosure placeholder) ----
    heading(doc, "", "Acknowledgment")
    body(doc, "[Disclosure placeholder — per the venue's policy, state any use of AI writing "
              "assistance here, e.g., grammar/formatting support, with the authors accountable for all "
              "content. Remove or adapt as required.]", indent=0)

    # ---- References ----
    heading(doc, "", "References")
    refs = [
        "K. K. Wong et al., “Fluid antenna systems,” IEEE Trans. Wireless Commun., 2021. [VERIFY]",
        "Switching-cost-aware deep reinforcement learning for dynamic port selection in FAS, IEEE "
        "Commun. Lett., 2026. [VERIFY]",
        "J. Zou, S. Sun, and C. Wang, “Online learning-induced port selection for fluid antenna in "
        "dynamic channel environment,” IEEE Wireless Commun. Lett., 2024. [VERIFY]",
        "Z. Zhang, J. Zhu, L. Dai, and R. W. Heath, “Successive Bayesian reconstructor for channel "
        "estimation in fluid antenna systems,” IEEE Trans. Wireless Commun., 2025. [VERIFY]",
        "C. Skouroumounis and I. Krikidis, “Fluid antenna with linear MMSE channel estimation for "
        "large-scale cellular networks,” IEEE Trans. Commun., 2023. [VERIFY]",
        "T. Parr, G. Pezzulo, and K. J. Friston, Active Inference: The Free Energy Principle in Mind, "
        "Brain, and Behavior. MIT Press, 2022. [VERIFY]",
    ]
    for i, rf in enumerate(refs, 1):
        p = doc.add_paragraph(); p.paragraph_format.left_indent = Inches(0.18)
        p.paragraph_format.first_line_indent = Inches(-0.18); p.paragraph_format.space_after = Pt(1)
        r = p.add_run(f"[{i}] {rf}"); r.font.name = TNR; r.font.size = Pt(8)

    doc.save(OUT)
    print("saved", OUT)


if __name__ == "__main__":
    build()
