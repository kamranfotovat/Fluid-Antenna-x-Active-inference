"""
Step-by-step IEEE working draft (grows one section at a time).

This is the TEMP working version we build with the user, section by section, so
we can annotate it by hand between steps. It deliberately contains ONLY the parts
we have agreed on so far -- right now: Title, Authors, Abstract, Index Terms, and
Section I (Introduction), plus the references the Introduction cites.

No figures and no Sections II-V yet: those arrive in later steps. The completed
full draft lives separately in make_paper_docx.py / Active_Inference_FAS_draft.docx.

IEEE fidelity: Times New Roman; 24pt centered title; author block; single-column
abstract/index terms; continuous section break into a two-column body; centered
small-caps section headings with roman numerals.

Output: paper/working_draft.docx
Run:    python paper/make_working_draft.py
"""

from __future__ import annotations

import os
import tempfile
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
from docx.enum.section import WD_SECTION
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

import matplotlib
matplotlib.use("Agg")
import matplotlib.pyplot as plt
matplotlib.rcParams["mathtext.fontset"] = "stix"   # Times-like math to match the body
matplotlib.rcParams["font.family"] = "STIXGeneral"
try:
    from PIL import Image
except Exception:  # pragma: no cover
    Image = None

HERE = os.path.dirname(os.path.abspath(__file__))
OUT = os.path.join(HERE, "working_draft.docx")

TNR = "Times New Roman"

# ---- IEEE-style equation rendering (LaTeX-quality math as embedded images) ----
# python-docx writes plain Unicode, which does not look like IEEE math. Instead we
# render each relation with matplotlib mathtext to a tight transparent PNG, then embed
# it centered with a right-flush equation number, mimicking the IEEE two-column look.
_EQDIR = tempfile.mkdtemp(prefix="fas_eq_")
_EQCNT = {"n": 0}
COL_W = 3.45   # ~two-column text width (in)


def render_eq(latex, fontsize=14, dpi=300):
    _EQCNT["n"] += 1
    path = os.path.join(_EQDIR, f"eq_{_EQCNT['n']}.png")
    fig = plt.figure(figsize=(0.1, 0.1))
    fig.text(0.5, 0.5, f"${latex}$", ha="center", va="center", fontsize=fontsize)
    fig.savefig(path, dpi=dpi, transparent=True, bbox_inches="tight", pad_inches=0.04)
    plt.close(fig)
    return path, dpi


def equation_img(doc, latex, num, maxw=2.85):
    path, dpi = render_eq(latex)
    if Image is not None:
        pw, ph = Image.open(path).size
        w_in, h_in = pw / dpi, ph / dpi
    else:
        w_in, h_in = maxw, 0.3
    if w_in > maxw:
        s = maxw / w_in; w_in *= s; h_in *= s
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(5); pf.space_after = Pt(5); pf.first_line_indent = Inches(0)
    pf.tab_stops.add_tab_stop(Inches(COL_W / 2.0), WD_TAB_ALIGNMENT.CENTER)
    pf.tab_stops.add_tab_stop(Inches(COL_W), WD_TAB_ALIGNMENT.RIGHT)
    p.add_run("\t")
    p.add_run().add_picture(path, width=Inches(w_in))
    rn = p.add_run(f"\t({num})"); rn.font.name = TNR; rn.font.size = Pt(10)
    return p


def set_cols(section, n):
    sectPr = section._sectPr
    cols = sectPr.find(qn("w:cols"))
    if cols is None:
        cols = OxmlElement("w:cols"); sectPr.append(cols)
    cols.set(qn("w:num"), str(n))
    cols.set(qn("w:space"), "360")   # ~0.25in gutter


def base_font(doc):
    st = doc.styles["Normal"]
    st.font.name = TNR
    st.font.size = Pt(10)
    st.element.rPr.rFonts.set(qn("w:eastAsia"), TNR)
    pf = st.paragraph_format
    pf.space_after = Pt(0); pf.line_spacing = 1.0


def para(doc, text="", size=10, bold=False, italic=False, align=None, after=0, before=0,
         first_indent=0.2):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    pf = p.paragraph_format
    pf.space_after = Pt(after); pf.space_before = Pt(before)
    if first_indent:
        pf.first_line_indent = Inches(first_indent)
    r = p.add_run(text)
    r.font.name = TNR; r.font.size = Pt(size); r.bold = bold; r.italic = italic
    return p, r


def heading(doc, num, title):
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(8); p.paragraph_format.space_after = Pt(4)
    label = f"{num}. {title.upper()}" if num else title.upper()
    r = p.add_run(label)
    r.font.name = TNR; r.font.size = Pt(10); r.font.small_caps = True; r.bold = False
    return p


def _pbdr(p, top=False, bottom=False, sz=8):
    """Add top/bottom rules to a paragraph (for the IEEE algorithm box)."""
    pPr = p._p.get_or_add_pPr()
    pbdr = OxmlElement("w:pBdr")
    for side, on in (("top", top), ("bottom", bottom)):
        if on:
            e = OxmlElement(f"w:{side}")
            e.set(qn("w:val"), "single"); e.set(qn("w:sz"), str(sz))
            e.set(qn("w:space"), "3"); e.set(qn("w:color"), "000000")
            pbdr.append(e)
    pPr.append(pbdr)


def algorithm(doc, number, title, lines, fs=8.5):
    """IEEE-style boxed algorithm. lines: list of (num|None, indent, [(text,bold),...])."""
    cap = doc.add_paragraph()
    pf = cap.paragraph_format
    pf.space_before = Pt(6); pf.space_after = Pt(2); pf.first_line_indent = Inches(0)
    _pbdr(cap, top=True, bottom=True, sz=10)
    r = cap.add_run(f"Algorithm {number}  "); r.font.name = TNR; r.font.size = Pt(fs + 0.5); r.bold = True
    r2 = cap.add_run(title); r2.font.name = TNR; r2.font.size = Pt(fs + 0.5); r2.bold = True
    for num, ind, segs in lines:
        p = doc.add_paragraph()
        pf = p.paragraph_format
        pf.space_before = Pt(0); pf.space_after = Pt(0)
        pf.left_indent = Inches(0.22); pf.first_line_indent = Inches(-0.22)  # hanging under the number
        lead = f"{num}: " if num is not None else ""
        r = p.add_run(lead); r.font.name = TNR; r.font.size = Pt(fs)
        if ind:
            s = p.add_run(" " * ind); s.font.name = TNR; s.font.size = Pt(fs)
        for txt, bold in segs:
            rr = p.add_run(txt); rr.font.name = TNR; rr.font.size = Pt(fs); rr.bold = bold
    close = doc.add_paragraph()
    close.paragraph_format.space_before = Pt(0); close.paragraph_format.space_after = Pt(5)
    _pbdr(close, top=True, sz=10)
    rc = close.add_run(""); rc.font.size = Pt(2)


def subhead(doc, letter, title):
    p = doc.add_paragraph(); p.paragraph_format.space_before = Pt(5); p.paragraph_format.space_after = Pt(1)
    r = p.add_run(f"{letter}. {title}")
    r.font.name = TNR; r.font.size = Pt(10); r.italic = True
    return p


def body(doc, text, indent=0.2):
    return para(doc, text, size=10, align=WD_ALIGN_PARAGRAPH.JUSTIFY, first_indent=indent)


def build():
    doc = Document()
    base_font(doc)
    sec = doc.sections[0]
    sec.top_margin = Inches(0.75); sec.bottom_margin = Inches(1.0)
    sec.left_margin = Inches(0.62); sec.right_margin = Inches(0.62)
    set_cols(sec, 1)

    # ---- Title + authors (single column) ----
    para(doc, "Active Inference for Port Selection in Fluid Antenna Systems Under Partial CSI",
         size=24, bold=False, align=WD_ALIGN_PARAGRAPH.CENTER, after=10, first_indent=0)
    para(doc, "Kian Fotovat¹ and Kamran Fotovat²", size=11,
         align=WD_ALIGN_PARAGRAPH.CENTER, after=2, first_indent=0)
    para(doc, "¹Dept. of Electrical and Computer Engineering, University of Tehran, Tehran, Iran",
         size=9, italic=True, align=WD_ALIGN_PARAGRAPH.CENTER, after=0, first_indent=0)
    para(doc, "²Iran University of Science and Technology, Tehran, Iran",
         size=9, italic=True, align=WD_ALIGN_PARAGRAPH.CENTER, after=10, first_indent=0)

    # ---- Abstract (single column) ----
    ab = doc.add_paragraph(); ab.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    ab.paragraph_format.left_indent = Inches(0.3); ab.paragraph_format.right_indent = Inches(0.3)
    r = ab.add_run("Abstract—"); r.font.name = TNR; r.font.size = Pt(9); r.bold = True; r.italic = True
    r2 = ab.add_run(
        "Fluid antenna systems (FAS) enhance spatial diversity by activating a few radiating ports "
        "among many candidate positions, but selecting good ports conventionally requires channel state "
        "information (CSI) across all ports, incurring prohibitive acquisition overhead. We recast FAS "
        "port selection as active inference under partial CSI: the transmitter measures only the ports it "
        "activates, infers the remaining ports from a spatio-temporal generative model (Jakes spatial "
        "correlation and a first-order temporal model) through a complex Kalman belief, and selects ports "
        "by minimizing the expected free energy (EFE), which unifies expected achievable rate, information "
        "gain, and antenna-switching cost in a single objective. Selection is beamforming-aware: its "
        "pragmatic term is the expected sum rate of a robust minimum-mean-square-error (MMSE) precoder "
        "built from the same belief, which then serves the users. A greedy submodular selector yields low "
        "decision latency, and an observe-then-precode protocol keeps the served channel fresh. "
        "Simulations show that the proposed method operates on a rate–switching frontier reaching up to "
        "89% of a full-CSI genie's sum rate while observing only 20% of the ports; its switching-aware "
        "objective exceeds that of the genie across the entire frontier and, at a balanced operating "
        "point, surpasses naive, deep-reinforcement-learning, and bandit baselines. The agent requires no "
        "training and remains robust to Doppler and to model mismatch, for which the spatial correlation "
        "can be learned online.")
    r2.font.name = TNR; r2.font.size = Pt(9); r2.italic = True

    # ---- Index terms (single column) ----
    it = doc.add_paragraph(); it.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    it.paragraph_format.left_indent = Inches(0.3); it.paragraph_format.right_indent = Inches(0.3)
    it.paragraph_format.space_before = Pt(4); it.paragraph_format.space_after = Pt(6)
    r = it.add_run("Index Terms—"); r.font.name = TNR; r.font.size = Pt(9); r.bold = True; r.italic = True
    r2 = it.add_run("Fluid antenna systems, port selection, active inference, expected free energy, "
                    "partial CSI, Kalman filtering.")
    r2.font.name = TNR; r2.font.size = Pt(9); r2.italic = True

    # ---- switch to two columns ----
    doc.add_section(WD_SECTION.CONTINUOUS)
    set_cols(doc.sections[-1], 2)

    # ===================== I. Introduction =====================
    heading(doc, "I", "Introduction")
    body(doc,
         "Fluid antenna systems (FAS), also known as movable- or reconfigurable-position antenna systems, "
         "replace fixed radiating elements with elements whose location can be adjusted on demand among a "
         "dense set of candidate positions, or ports, within a compact aperture [1]. By repositioning to a "
         "favorable point in the spatial fading pattern, an FAS harvests spatial diversity that a "
         "conventional fixed array of the same size cannot, and it does so with only a handful of "
         "radio-frequency (RF) chains, since at any instant only the activated ports are connected to the "
         "front end. The ability to place antennas where the channel is strong, rather than where a fixed "
         "grid happens to lie, has made FAS a promising ingredient for next-generation wireless systems, "
         "with reported gains in diversity, spatial multiplexing, and multiple access. Across all of these "
         "settings the pivotal control problem is the same: port selection, that is, deciding slot by slot "
         "which subset of the many candidate ports to activate so as to maximize the delivered "
         "throughput.", indent=0)
    body(doc,
         "The difficulty is that a port can be scored only if its channel is known, and the channel "
         "differs from port to port. Conventional selection therefore presupposes channel state "
         "information (CSI) across all candidate ports. Acquiring that CSI, however, is expensive: the "
         "ports are numerous and, because the aperture is sampled at sub-wavelength spacing, they are "
         "packed far more densely than a half-wavelength array, so the number of channel coefficients to "
         "be estimated, and the pilot resources they consume, grows rapidly with the aperture. Moreover, "
         "only the activated ports can actually be observed through the RF chains; the remaining ports are "
         "at any instant physically unmeasured. The true bottleneck of FAS is thus not the selection "
         "arithmetic but channel acquisition under a hard sensing budget: the system must decide not only "
         "which ports to serve, but which ports are even worth measuring. This sensing cost is precisely "
         "what existing formulations tend to assume away.")
    body(doc,
         "Two broad lines of work approach the problem but leave this gap open. The first is learning-based "
         "selection. Deep reinforcement learning (DRL) with attention over ports [2] and online bandit "
         "methods [3] learn selection policies directly, and recent switching-cost-aware variants even "
         "penalize antenna movement; these methods, however, typically treat the full channel as an "
         "observed, cost-free input to the policy and/or require extensive interaction to train. The "
         "second is channel estimation for FAS. The successive Bayesian reconstructor [4] and sequential "
         "linear minimum-mean-square-error estimation [5] exploit spatial correlation to reconstruct the "
         "channel from limited pilots with high accuracy. Yet these methods solve estimation in isolation: "
         "they recover the channel but do not decide what to sense, what to serve, or when to move, and "
         "they do not weigh the value of an additional measurement against its cost. What is missing is a "
         "formulation in which sensing, serving, and movement are chosen jointly, under an explicit "
         "budget, from partial observations.")
    body(doc,
         "We contend that active inference supplies exactly this missing perspective. Rooted in the "
         "free-energy principle from computational neuroscience [6], active inference casts a "
         "decision-maker as an agent that maintains a probabilistic generative model of its environment "
         "and selects actions that minimize an expected free energy (EFE). A distinctive feature of the "
         "EFE is that it decomposes into two terms with clear operational meaning: a pragmatic value that "
         "rewards actions expected to fulfil the agent's goals, and an epistemic value that rewards "
         "actions expected to be informative, that is, to reduce uncertainty about hidden states of the "
         "world. Perception, the updating of beliefs from observations, and action, the choice of what to "
         "do and thereby what to observe next, are handled within one coherent objective rather than as "
         "separate mechanisms. This is a strikingly natural fit for FAS port selection under partial CSI. "
         "The hidden state is the channel over the unmeasured ports; the generative model is the spatial "
         "and temporal correlation structure of that channel; the pragmatic value is the throughput a port "
         "set is expected to deliver; and the epistemic value is the information a measurement reveals "
         "about the ports left unseen. The very tension that defines the problem, namely whether to spend "
         "a scarce activated port to serve now or to probe an uncertain but promising region for later, is "
         "expressed directly by the EFE rather than engineered in by hand. In contrast to black-box "
         "learning, the resulting agent is model-based and interpretable: every decision traces back to a "
         "belief about the channel and a transparent value of information, and it needs no training data. "
         "To the best of our knowledge, active inference has not previously been applied to fluid-antenna "
         "port selection, nor to channel acquisition under a sensing budget more broadly.")
    body(doc,
         "Building on this view, we formulate FAS port selection as active inference under partial CSI and "
         "develop a concrete, low-complexity agent for it. Our contributions are as follows:")
    body(doc,
         "1) We formulate FAS port selection as active inference under partial observability, removing "
         "the unrealistic full-CSI assumption of prior work and casting the task as active channel "
         "acquisition under a sensing budget.", indent=0)
    body(doc,
         "2) We develop a model-based agent that maintains a complex Kalman belief over all ports from a "
         "spatio-temporal prior, and selects ports by minimizing an expected free energy that unifies "
         "expected rate, information gain, and switching cost in a single objective, with no separately "
         "tuned preference vector. Selection is coupled to transmission: its pragmatic value is the "
         "achievable sum rate of a robust MMSE precoder computed from the belief, so ports are chosen in "
         "anticipation of the beamformer that will serve them.", indent=0)
    body(doc,
         "3) We give a greedy submodular selector with an (1−1/e) guarantee on the information term "
         "and O(NM) per-slot latency, together with an observe-then-precode protocol that keeps the served "
         "channel fresh and makes performance robust to Doppler.", indent=0)
    body(doc,
         "4) Through simulation we show that the agent operates on a rate–switching frontier reaching up "
         "to 89% of a full-CSI genie's sum rate while observing only 20% of the ports; its switching-aware "
         "objective exceeds the genie's across the whole frontier and, at a balanced operating point, "
         "surpasses naive, DRL, and bandit baselines with zero training, while remaining robust to Doppler "
         "and able to learn the spatial correlation online under model mismatch.", indent=0)

    # ===================== II. Fluid Antenna System Model =====================
    heading(doc, "II", "Fluid Antenna System Model")
    body(doc,
         "We consider the downlink of a single-cell multiuser system in which a base station (BS) equipped "
         "with a fluid antenna serves K single-antenna users. This section specifies the array geometry, "
         "the spatio-temporal channel, the partial observation model that is central to this work, the "
         "transmission and rate model, and the switching-aware objective.", indent=0)

    subhead(doc, "A", "Array Geometry and Activation")
    body(doc,
         "The BS fluid antenna comprises N = Nₓ×Nᵧ candidate ports arranged on a uniform planar grid over "
         "a compact aperture with sub-half-wavelength spacing. The BS has M radio-frequency (RF) chains, "
         "M ≤ N, so in each slot t it activates a subset S(t) ⊆ {1,…,N} of |S(t)| = M ports; only "
         "activated ports radiate and connect to the front end. Serving K streams requires M ≥ K, hence "
         "N ≥ M ≥ K. We write P_S ∈ {0,1}^{M×N} for the selection matrix whose rows are the standard basis "
         "vectors indexed by S, so that P_S x extracts the entries of a vector x on the activated ports.",
         indent=0)

    subhead(doc, "B", "Spatio-Temporal Channel Model")
    body(doc,
         "Let hₖ(t) ∈ ℂᴺ denote the channel between the N candidate ports and user k in slot t. Because "
         "the ports are densely spaced, their coefficients are spatially correlated. Following the classical "
         "Jakes model for rich scattering, the correlation between ports i and j depends only on their "
         "separation dᵢⱼ,", indent=0)
    equation_img(doc, r"R_{ij} = J_0\left(\frac{2\pi d_{ij}}{\lambda}\right)", 1)
    body(doc,
         "where J₀(·) is the zeroth-order Bessel function of the first kind and λ the wavelength. "
         "Collecting these entries into R ∈ ℂᴺˣᴺ, the small-scale fading is spatially correlated Rayleigh, "
         "hₖ ~ 𝒞𝒩(0, βₖR), with βₖ the large-scale gain of user k. The channel ages between slots "
         "according to a first-order Gauss–Markov (autoregressive) process,")
    equation_img(doc,
        r"\mathbf{h}_k(t) = \rho\,\mathbf{h}_k(t-1) + \sqrt{1-\rho^2}\,\mathbf{e}_k(t),\quad "
        r"\mathbf{e}_k(t)\sim\mathcal{CN}(\mathbf{0},\beta_k\mathbf{R})", 2)
    body(doc,
         "where eₖ(t) is independent of hₖ(t−1), and the temporal correlation ρ = J₀(2π f_D T_s) is set by "
         "the maximum Doppler frequency f_D and the slot duration T_s; ρ → 1 for slow fading and decreases "
         "with mobility. The process is stationary with marginal hₖ(t) ~ 𝒞𝒩(0, βₖR).")

    subhead(doc, "C", "Observation Model: Partial CSI")
    body(doc,
         "A defining constraint of our setting is that the BS cannot observe the full channel: information "
         "about a port is obtained only when that port is activated. During the pilot phase of slot t the "
         "users send orthogonal pilots and the BS estimates the channel on the M activated ports, while "
         "the remaining N−M ports are unobserved. The measurement is", indent=0)
    equation_img(doc,
        r"\mathbf{y}_k(t) = \mathbf{P}_{\mathcal{S}(t)}\,\mathbf{h}_k(t) + \mathbf{n}_k(t),\quad "
        r"\mathbf{n}_k(t)\sim\mathcal{CN}(\mathbf{0},\sigma_e^2\mathbf{I}_M)", 3)
    body(doc,
         "where P_{S(t)} restricts hₖ(t) to the activated ports and nₖ(t) is the pilot/estimation noise "
         "with per-port variance σ_e² set by the pilot SNR. Equivalently, the BS observes the length-M "
         "vector h_{k,S}(t) = P_{S(t)} hₖ(t) in noise and must infer the N−M hidden ports from spatial and "
         "temporal correlation. This partial observability distinguishes our formulation from prior "
         "port-selection work that assumes full-channel knowledge, and it is what makes the choice of "
         "which ports to measure consequential.")

    subhead(doc, "D", "Downlink Transmission and Achievable Rate")
    body(doc,
         "Over the activated ports the BS transmits with a linear precoder W(t) = [w₁(t),…,w_K(t)] ∈ "
         "ℂᴹˣᴷ, one column per user, under a total power budget ‖W(t)‖_F² ≤ P. The transmitted signal is "
         "x(t) = Σₖ wₖ(t) sₖ(t) with data symbols sₖ ~ 𝒞𝒩(0,1). User k receives", indent=0)
    equation_img(doc,
        r"r_k = \mathbf{h}_{k,\mathcal{S}}^{\mathrm{H}}\mathbf{w}_k s_k + "
        r"\sum_{j\neq k}\mathbf{h}_{k,\mathcal{S}}^{\mathrm{H}}\mathbf{w}_j s_j + z_k", 4)
    body(doc,
         "where h_{k,S} = P_S hₖ is the channel on the activated ports and zₖ ~ 𝒞𝒩(0, σ²) is receiver "
         "noise. Treating inter-user interference as noise, the SINR and achievable rate of user k are")
    equation_img(doc,
        r"\mathrm{SINR}_k = \frac{|\mathbf{h}_{k,\mathcal{S}}^{\mathrm{H}}\mathbf{w}_k|^2}"
        r"{\sum_{j\neq k}|\mathbf{h}_{k,\mathcal{S}}^{\mathrm{H}}\mathbf{w}_j|^2 + \sigma^2}", 5)
    equation_img(doc, r"R_k = \log_2\left(1 + \mathrm{SINR}_k\right)", 6)
    body(doc,
         "The precoder is obtained by regularized MMSE (transmit-Wiener) filtering from the available "
         "channel estimate Ĥ = [ĥ_{1,S},…,ĥ_{K,S}] ∈ ℂᴹˣᴷ,")
    equation_img(doc,
        r"\mathbf{W} = \xi\,\hat{\mathbf{H}}\left(\hat{\mathbf{H}}^{\mathrm{H}}\hat{\mathbf{H}} + "
        r"\frac{K\sigma^2}{P}\mathbf{I}_K\right)^{-1}", 7)
    body(doc,
         "where ξ scales W to meet the power budget. Under full CSI, Ĥ is the true activated channel; in "
         "our partial-CSI setting Ĥ and its error statistics are supplied by the belief of Section III, "
         "yielding a robust MMSE precoder that additionally accounts for the residual channel uncertainty.")

    subhead(doc, "E", "Switching Cost and Problem Formulation")
    body(doc,
         "Reconfiguring the fluid antenna between slots is not free: physically moving or re-tuning "
         "elements incurs delay and energy. We model this by the number of ports that change between "
         "consecutive slots,", indent=0)
    equation_img(doc,
        r"C_t = |\mathcal{S}_t\,\triangle\,\mathcal{S}_{t-1}|,\qquad "
        r"E_{\mathrm{sw}}(t) = e_{\mathrm{sw}}\,C_t", 8)
    body(doc,
         "where △ is the symmetric set difference and e_sw the per-port switching energy. The design goal "
         "is to choose the activation sequence that maximizes long-term throughput net of switching cost,")
    equation_img(doc,
        r"\max_{\{\mathcal{S}_t\}}\ \sum_{t=1}^{T}\left(\sum_{k=1}^{K} R_k(t) - "
        r"\eta_{\mathrm{sw}}\,E_{\mathrm{sw}}(t)\right)\ \ \mathrm{s.t.}\ |\mathcal{S}_t| = M", 9)
    body(doc,
         "where η_sw ≥ 0 trades throughput against reconfiguration cost. Equation (9) is the objective "
         "against which every method in this paper is evaluated. Its difficulty is that Rₖ(t) depends on "
         "the channel over ports that, under partial CSI, are largely unobserved — the problem we address "
         "next through active inference.")

    # ===================== III. Active Inference for Port Selection =====================
    heading(doc, "III", "Active Inference for Port Selection")
    body(doc,
         "We now cast port selection as active inference. The agent holds a generative model of the "
         "channel, maintains a Bayesian belief over all ports (perception), and each slot chooses the port "
         "set that minimizes an expected free energy (action), which trades expected rate against "
         "information gain and switching cost within one objective.", indent=0)

    subhead(doc, "A", "Generative Model")
    body(doc,
         "The latent state is the full channel {hₖ(t)}, of which only the activated ports are observed. "
         "The agent's generative model factorizes over users and time as", indent=0)
    equation_img(doc,
        r"p(\{\mathbf{h}_k,\mathbf{y}_k\}) = \prod_{k}p(\mathbf{h}_k(0))\prod_{t}"
        r"p(\mathbf{h}_k(t)\,|\,\mathbf{h}_k(t-1))\,p(\mathbf{y}_k(t)\,|\,\mathbf{h}_k(t),\mathcal{S}_t)", 10)
    body(doc,
         "with three ingredients already fixed by the physics of Section II: the prior p(hₖ(0)) = "
         "𝒞𝒩(0, βₖR); the transition p(hₖ(t)|hₖ(t−1)) given by the AR(1) dynamics (2); and the likelihood "
         "p(yₖ(t)|hₖ(t), Sₜ) given by the partial observation (3). Preferences complete the model. Unlike "
         "textbook active inference, in which preferred outcomes are encoded in a hand-tuned vector, here "
         "the preference is the communication utility itself — the achievable rate net of switching in "
         "(9). There is thus no arbitrary preference to set: the agent's goal is the system objective.")

    subhead(doc, "B", "Perception: Complex Kalman Belief")
    body(doc,
         "Because the model is linear-Gaussian, the exact posterior over each user's channel is Gaussian, "
         "q(hₖ) = 𝒞𝒩(μₖ, Σₖ), maintained by a per-user complex Kalman filter with two steps per slot. The "
         "predict step propagates the belief through the AR(1) dynamics, encoding channel aging,", indent=0)
    equation_img(doc,
        r"\mu_k \leftarrow \rho\,\mu_k,\qquad \Sigma_k \leftarrow \rho^2\Sigma_k + (1-\rho^2)\beta_k\mathbf{R}", 11)
    body(doc,
         "When the ports in Sₜ are activated and yₖ observed, the measurement update corrects the belief. "
         "With Kalman gain")
    equation_img(doc,
        r"\mathbf{K}_k = \Sigma_k\mathbf{P}_{\mathcal{S}}^{\mathrm{H}}\left(\mathbf{P}_{\mathcal{S}}"
        r"\Sigma_k\mathbf{P}_{\mathcal{S}}^{\mathrm{H}} + \sigma_e^2\mathbf{I}_M\right)^{-1}", 12)
    body(doc, "the posterior mean and covariance become")
    equation_img(doc,
        r"\mu_k \leftarrow \mu_k + \mathbf{K}_k\left(\mathbf{y}_k - \mathbf{P}_{\mathcal{S}}\mu_k\right)", 13)
    equation_img(doc,
        r"\Sigma_k \leftarrow \left(\mathbf{I} - \mathbf{K}_k\mathbf{P}_{\mathcal{S}}\right)\Sigma_k"
        r"\left(\mathbf{I} - \mathbf{K}_k\mathbf{P}_{\mathcal{S}}\right)^{\mathrm{H}} + "
        r"\sigma_e^2\mathbf{K}_k\mathbf{K}_k^{\mathrm{H}}", 14)
    body(doc,
         "where the Joseph form (14) keeps Σₖ symmetric positive semidefinite. Only the M×M innovation "
         "covariance is inverted, and it is regularized by σ_e²I; consequently the rank-deficient spatial "
         "prior R needs no artificial jitter. Crucially, observing a port also shrinks the uncertainty of "
         "its correlated neighbours through the off-diagonals of R and Σₖ, so a few measurements inform "
         "many ports. The filter is calibrated: its predicted covariance matches the realized estimation "
         "error.")

    subhead(doc, "C", "Action: Expected Free Energy")
    body(doc,
         "Each slot the agent scores a candidate port set S by its expected free energy, written here as a "
         "value to be minimized,", indent=0)
    equation_img(doc,
        r"G(\mathcal{S}) = -\,\mathrm{Prag}(\mathcal{S}) - \kappa\,\mathrm{Epis}(\mathcal{S}) + "
        r"\eta_{\mathrm{sw}}E_{\mathrm{sw}}(\mathcal{S})", 15)
    body(doc,
         "a pragmatic (goal) term, an epistemic (information) term weighted by an exploration weight κ ≥ 0, "
         "and the switching cost of (8). The pragmatic value is the expected sum rate the belief predicts "
         "for S under the robust MMSE precoder; accounting for residual CSI error through the belief "
         "covariance gives the imperfect-CSI lower bound")
    equation_img(doc,
        r"\mathrm{Prag}(\mathcal{S}) = \sum_{k}\log_2\left(1 + "
        r"\frac{|\mu_{k,\mathcal{S}}^{\mathrm{H}}\mathbf{w}_k|^2}"
        r"{\sum_{j\neq k}|\mu_{k,\mathcal{S}}^{\mathrm{H}}\mathbf{w}_j|^2 + "
        r"\sum_j \mathbf{w}_j^{\mathrm{H}}\Sigma_{k,\mathcal{S}}\mathbf{w}_j + \sigma^2}\right)", 16)
    body(doc,
         "where μ_{k,S} = P_S μₖ, Σ_{k,S} = P_S Σₖ P_Sᴴ, and the error term Σⱼ wⱼᴴ Σ_{k,S} wⱼ makes the "
         "agent conservative when uncertain; the robust MMSE precoder augments (7) with Σ_{k,S}. The "
         "epistemic value is the mutual information between the channel and the pilots that S would yield, "
         "i.e., the expected reduction in belief entropy,")
    equation_img(doc,
        r"\mathrm{Epis}(\mathcal{S}) = \sum_{k}\log_2\det\left(\mathbf{I}_M + "
        r"\frac{\mathbf{P}_{\mathcal{S}}\Sigma_k\mathbf{P}_{\mathcal{S}}^{\mathrm{H}}}{\sigma_e^2}\right)", 17)
    body(doc,
         "which is monotone and submodular in S; through the correlation in Σₖ, activating one port also "
         "gathers information about others. The two terms embody the exploit–explore tension natively: "
         "the pragmatic value favours ports that serve well now, the epistemic value favours ports that "
         "most reduce uncertainty for later, and κ sets the balance — no separate exploration heuristic is "
         "bolted on.")

    subhead(doc, "D", "Observe-Then-Precode Protocol")
    body(doc,
         "The order of operations within a slot matters. We adopt an observe-then-precode protocol: "
         "(i) predict the belief; (ii) select Sₜ by minimizing G on the predicted belief; (iii) activate "
         "Sₜ, obtain the pilots yₖ, and update the belief; and (iv) form the robust MMSE precoder from the "
         "updated belief and transmit. Because the served ports are re-measured before precoding, their "
         "posterior error is about σ_e² rather than the one-step aging floor (1−ρ²)βₖ that a "
         "predict-then-act ordering incurs. As shown in Section IV, this makes the served rate nearly "
         "independent of Doppler.", indent=0)

    subhead(doc, "E", "Greedy Submodular Selection")
    body(doc,
         "Selecting the best M-of-N set exactly is combinatorial. We instead build Sₜ greedily: starting "
         "from the empty set, we repeatedly add the port whose inclusion most decreases G,", indent=0)
    equation_img(doc,
        r"p^{\star} = \arg\max_{p\notin\mathcal{S}}\ \left[\,G(\mathcal{S}) - "
        r"G(\mathcal{S}\cup\{p\})\,\right]", 18)
    body(doc,
         "until |S| = M. Since the epistemic term is monotone submodular and the switching term modular, "
         "greedy selection attains the standard (1−1/e) near-optimality guarantee on the information-driven "
         "part, at O(NM) score evaluations per slot — about 20 ms in our setting versus seconds for "
         "exhaustive search, a roughly 500× reduction. Algorithm 1 summarizes the complete per-slot agent, "
         "with the greedy loop (lines 3–8) embedded in the observe-then-precode ordering of Section III-D.")

    B = True; N = False  # bold / not-bold shorthand for the pseudocode below
    algorithm(doc, 1, "Active-Inference Port Selection (slot t)", [
        (None, 0, [("Require: ", B), ("predicted belief {μₖ, Σₖ}, previous set Sₜ₋₁, budget M, weight κ", N)]),
        (None, 0, [("Ensure: ", B), ("activated set Sₜ and precoder W", N)]),
        (1, 0, [("predict belief: μₖ ← ρμₖ,  Σₖ ← ρ²Σₖ + (1−ρ²)βₖR   ∀k", N)]),
        (2, 0, [("S ← ∅", N)]),
        (3, 0, [("for", B), (" m = 1 to M ", N), ("do", B)]),
        (4, 2, [("for each", B), (" candidate port p ∉ S ", N), ("do", B)]),
        (5, 4, [("Δ(p) ← G(S) − G(S ∪ {p})", N), ("      ▷ via (15)–(17)", N)]),
        (6, 2, [("end for", B)]),
        (7, 2, [("p* ← arg maxₚ Δ(p);   S ← S ∪ {p*}", N)]),
        (8, 0, [("end for", B)]),
        (9, 0, [("Sₜ ← S", N)]),
        (10, 0, [("activate Sₜ, observe pilots yₖ, update belief   ▷ via (12)–(14)", N)]),
        (11, 0, [("W ← robust MMSE precoder from updated {μₖ, Σₖ}", N)]),
        (12, 0, [("transmit; incur switching cost η_sw · |Sₜ △ Sₜ₋₁|", N)]),
    ])

    subhead(doc, "F", "Online Learning of the Spatial Correlation")
    body(doc,
         "The agent assumes a Jakes R, but real propagation may differ. When R is unknown or non-Jakes it "
         "can be learned online from the accumulated pilots: averaging the outer products of co-observed "
         "pilots and removing the noise floor gives", indent=0)
    equation_img(doc,
        r"\hat{R}_{ij} = \frac{\sum_t \mathbf{1}\{i,j\in\mathcal{S}_t\}\,y_i(t)\,y_j^{*}(t)}"
        r"{\sum_t \mathbf{1}\{i,j\in\mathcal{S}_t\}} - \sigma_e^2\,\delta_{ij}", 19)
    body(doc,
         "which is then normalized to unit diagonal (cancelling the per-user power βₖ) and projected onto "
         "the positive-semidefinite cone. Substituting R̂ for R restores performance under model mismatch "
         "(Section IV); in the matched case the belief is already Bayes-optimal, so learning can only help "
         "— consistent with the method's zero-training nature.")

    # ---- References cited by the Introduction so far ----
    heading(doc, "", "References")
    refs = [
        "K. K. Wong et al., “Fluid antenna systems,” IEEE Trans. Wireless Commun., 2021. [VERIFY]",
        "Switching-cost-aware deep reinforcement learning for dynamic port selection in FAS, IEEE "
        "Commun. Lett., 2026. [VERIFY]",
        "J. Zou, S. Sun, and C. Wang, “Online learning-induced port selection for fluid antenna in "
        "dynamic channel environment,” IEEE Wireless Commun. Lett., 2024. [VERIFY]",
        "Z. Zhang, J. Zhu, L. Dai, and R. W. Heath, “Successive Bayesian reconstructor for channel "
        "estimation in fluid antenna systems,” IEEE Trans. Wireless Commun., 2025. [VERIFY]",
        "C. Skouroumounis and I. Krikidis, “Fluid antenna with linear MMSE channel estimation for "
        "large-scale cellular networks,” IEEE Trans. Commun., 2023. [VERIFY]",
        "T. Parr, G. Pezzulo, and K. J. Friston, Active Inference: The Free Energy Principle in Mind, "
        "Brain, and Behavior. MIT Press, 2022. [VERIFY]",
    ]
    for i, rf in enumerate(refs, 1):
        p = doc.add_paragraph(); p.paragraph_format.left_indent = Inches(0.18)
        p.paragraph_format.first_line_indent = Inches(-0.18); p.paragraph_format.space_after = Pt(1)
        r = p.add_run(f"[{i}] {rf}"); r.font.name = TNR; r.font.size = Pt(8)

    doc.save(OUT)
    print("saved", OUT)


if __name__ == "__main__":
    build()
